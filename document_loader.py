"""
DocumentLoader — PDF / CSV / Image / Text 멀티모달 로더

다이어그램 v8: 상단 입력 레이어
  - 텍스트 질의
  - 파일 입력 (PDF·CSV·Img)
  - 멀티모달 (Claude/GPT Vision LLM)

지원 포맷:
  PDF  → pypdf (텍스트 추출) + base64 (Vision 멀티모달)
  CSV  → csv 표준 라이브러리 또는 pandas
  DOCX → python-docx로 문단 텍스트 추출
  Img  → base64 인코딩 → Vision LLM 전달용 content block
  Text → UTF-8 직접 읽기
"""

import base64
import csv
import io
import json
import mimetypes
import os
from dataclasses import dataclass, field
from pathlib import Path
from typing import Any


# ══════════════════════════════════════════════════════════════════
#  데이터 클래스
# ══════════════════════════════════════════════════════════════════

@dataclass
class Document:
    """단일 문서 단위."""
    content: str                          # 텍스트 내용
    metadata: dict[str, Any] = field(default_factory=dict)
    page_number: int = 0
    doc_type: str = "text"               # "text" | "pdf" | "csv" | "image"
    image_data: str | None = None        # base64 (이미지/PDF 페이지)
    image_mime: str | None = None        # e.g. "image/png"


@dataclass
class LoadResult:
    """로드 결과 모음."""
    documents: list[Document]
    source: str
    total_pages: int = 0
    error: str | None = None

    def to_text(self) -> str:
        """모든 문서를 단순 텍스트로 결합."""
        return "\n\n".join(
            f"[페이지 {d.page_number}]\n{d.content}" if d.page_number else d.content
            for d in self.documents
        )


# ══════════════════════════════════════════════════════════════════
#  PDF 로더
# ══════════════════════════════════════════════════════════════════

def load_pdf(path: str | Path) -> LoadResult:
    """
    PDF 파일에서 텍스트를 추출한다.
    pypdf가 없으면 안내 메시지를 반환한다.
    """
    path = Path(path)
    try:
        import pypdf  # optional dependency
    except ImportError:
        return LoadResult(
            documents=[Document(content="[Error] pypdf가 설치되지 않았습니다. pip install pypdf")],
            source=str(path),
            error="pypdf not installed",
        )

    docs: list[Document] = []
    try:
        reader = pypdf.PdfReader(str(path))
        for i, page in enumerate(reader.pages):
            text = page.extract_text() or ""
            docs.append(
                Document(
                    content=text.strip(),
                    metadata={"source": str(path), "page": i + 1},
                    page_number=i + 1,
                    doc_type="pdf",
                )
            )
        return LoadResult(documents=docs, source=str(path), total_pages=len(reader.pages))
    except Exception as exc:
        return LoadResult(
            documents=[],
            source=str(path),
            error=str(exc),
        )


# ══════════════════════════════════════════════════════════════════
#  CSV 로더
# ══════════════════════════════════════════════════════════════════

def load_csv(path: str | Path, as_markdown: bool = True) -> LoadResult:
    """
    CSV 파일을 마크다운 테이블 또는 JSON 형식으로 로드한다.
    """
    path = Path(path)
    rows: list[dict] = []
    try:
        with open(path, newline="", encoding="utf-8") as f:
            reader = csv.DictReader(f)
            for row in reader:
                rows.append(dict(row))
    except Exception as exc:
        return LoadResult(documents=[], source=str(path), error=str(exc))

    if not rows:
        return LoadResult(
            documents=[Document(content="(빈 CSV 파일)", metadata={"source": str(path)}, doc_type="csv")],
            source=str(path),
        )

    if as_markdown:
        headers = list(rows[0].keys())
        header_line = "| " + " | ".join(headers) + " |"
        sep_line = "| " + " | ".join(["---"] * len(headers)) + " |"
        data_lines = [
            "| " + " | ".join(str(row.get(h, "")) for h in headers) + " |"
            for row in rows
        ]
        content = "\n".join([header_line, sep_line] + data_lines)
    else:
        content = json.dumps(rows, ensure_ascii=False, indent=2)

    doc = Document(
        content=content,
        metadata={"source": str(path), "rows": len(rows)},
        doc_type="csv",
    )
    return LoadResult(documents=[doc], source=str(path), total_pages=1)


# ══════════════════════════════════════════════════════════════════
#  DOCX 로더
# ══════════════════════════════════════════════════════════════════

def load_docx(path: str | Path) -> LoadResult:
    """
    DOCX(Word) 파일에서 문단 텍스트를 추출한다.
    python-docx가 없으면 안내 메시지를 반환한다.
    """
    path = Path(path)
    try:
        from docx import Document as WordDocument  # optional dependency
    except ImportError:
        return LoadResult(
            documents=[Document(content="[Error] python-docx가 설치되지 않았습니다. pip install python-docx")],
            source=str(path),
            error="python-docx not installed",
        )

    try:
        doc = WordDocument(str(path))
        paragraphs = [p.text.strip() for p in doc.paragraphs if p.text and p.text.strip()]
        content = "\n".join(paragraphs).strip()
        if not content:
            content = "(DOCX에서 추출 가능한 텍스트가 없습니다.)"
        return LoadResult(
            documents=[
                Document(
                    content=content,
                    metadata={"source": str(path)},
                    doc_type="text",
                )
            ],
            source=str(path),
            total_pages=1,
        )
    except Exception as exc:
        return LoadResult(documents=[], source=str(path), error=str(exc))


# ══════════════════════════════════════════════════════════════════
#  Image 로더 (멀티모달용 base64 인코딩)
# ══════════════════════════════════════════════════════════════════

def load_image(path: str | Path) -> LoadResult:
    """
    이미지를 base64로 인코딩한다.
    반환된 Document.image_data를 Vision LLM content block에 사용한다.
    """
    path = Path(path)
    mime, _ = mimetypes.guess_type(str(path))
    if mime not in ("image/png", "image/jpeg", "image/gif", "image/webp"):
        mime = "image/png"  # 기본값

    try:
        with open(path, "rb") as f:
            raw = f.read()
        encoded = base64.standard_b64encode(raw).decode()
    except Exception as exc:
        return LoadResult(documents=[], source=str(path), error=str(exc))

    doc = Document(
        content=f"[이미지 파일: {path.name}] (Vision LLM으로 분석 필요)",
        metadata={"source": str(path), "mime": mime, "size_bytes": len(raw)},
        doc_type="image",
        image_data=encoded,
        image_mime=mime,
    )
    return LoadResult(documents=[doc], source=str(path), total_pages=1)


def image_to_vision_content(doc: Document) -> list[dict]:
    """
    Document(image)를 OpenAI/Claude Vision API content block으로 변환한다.

    사용 예:
        content = image_to_vision_content(doc)
        # → [{"type": "image_url", "image_url": {"url": "data:image/png;base64,..."}}]
    """
    if not doc.image_data:
        return [{"type": "text", "text": doc.content}]

    return [
        {
            "type": "image_url",
            "image_url": {
                "url": f"data:{doc.image_mime};base64,{doc.image_data}"
            },
        },
        {"type": "text", "text": "이 이미지의 내용을 분석해주세요."},
    ]


# ══════════════════════════════════════════════════════════════════
#  텍스트 로더
# ══════════════════════════════════════════════════════════════════

def load_text(path: str | Path) -> LoadResult:
    path = Path(path)
    try:
        content = path.read_text(encoding="utf-8")
    except Exception as exc:
        return LoadResult(documents=[], source=str(path), error=str(exc))

    doc = Document(
        content=content,
        metadata={"source": str(path)},
        doc_type="text",
    )
    return LoadResult(documents=[doc], source=str(path), total_pages=1)


# ══════════════════════════════════════════════════════════════════
#  범용 로더 (확장자 자동 감지)
# ══════════════════════════════════════════════════════════════════

def load_document(path: str | Path) -> LoadResult:
    """
    파일 확장자를 기반으로 적합한 로더를 자동으로 선택한다.

    지원:
      .pdf           → load_pdf
      .csv           → load_csv
      .docx          → load_docx
      .png .jpg .jpeg .gif .webp → load_image
      그 외           → load_text
    """
    path = Path(path)
    ext = path.suffix.lower()

    if ext == ".pdf":
        return load_pdf(path)
    elif ext == ".csv":
        return load_csv(path)
    elif ext == ".docx":
        return load_docx(path)
    elif ext in (".png", ".jpg", ".jpeg", ".gif", ".webp"):
        return load_image(path)
    else:
        return load_text(path)


def load_documents(paths: list[str | Path]) -> list[Document]:
    """여러 파일을 한꺼번에 로드해 Document 리스트로 반환한다."""
    docs: list[Document] = []
    for p in paths:
        result = load_document(p)
        if result.error:
            print(f"[DocumentLoader] 오류 ({p}): {result.error}")
        docs.extend(result.documents)
    return docs
